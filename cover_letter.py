"""Generate tailored cover letters from a job description."""
import argparse
import os
import sys
from datetime import date
from pathlib import Path
from cv_data import CVData, CV_DATA
from core.docx_builder import INK, MUTED, apply_font, setup_document_margins

SYSTEM_PROMPT = """You are an expert career coach specializing in software engineering roles.
Write a tailored, confident cover letter for the candidate below.
Rules:
- 3 paragraphs only: (1) why this role, (2) specific relevant experience, (3) closing
- No generic filler like \"I am excited to apply\"
- Never start a sentence with \"I\"
- Reference specific technologies from the job description
- Tone: professional but human, not robotic
- No date, address, or \"Dear Hiring Manager\" header — body only
- Max 250 words"""

def load_job_description(args: argparse.Namespace) -> str:
    if args.job and args.file: print('Error: Use either --job or --file, not both.', file=sys.stderr); sys.exit(1)
    try: text = args.job if args.job else Path(args.file).read_text(encoding='utf-8') if args.file else ''
    except OSError as error: print(f'Error: Could not read job description file: {error}', file=sys.stderr); sys.exit(1)
    text = text.strip()
    if len(text) < 30: print('Error: Job description must contain at least 30 characters.', file=sys.stderr); sys.exit(1)
    return text

def build_user_prompt(job_desc: str, data: CVData) -> str:
    contact = data.contact
    return f"""Candidate
Name: {contact.name}
Current title: {contact.title}
Location: {contact.location}
Skills: {data.skills_text}
Summary: {data.summary}
Experience: {'; '.join(f'{x.role} at {x.company}: {x.summary}' for x in data.experiences)}
Projects: {'; '.join(f'{x.title}: {x.description}' for x in data.projects)}

Job description
{job_desc}
"""

def call_anthropic(system: str, user: str) -> str:
    key = os.getenv('ANTHROPIC_API_KEY')
    if not key: print('Error: Set ANTHROPIC_API_KEY environment variable first.', file=sys.stderr); sys.exit(1)
    try:
        from anthropic import Anthropic
        response = Anthropic(api_key=key).messages.create(model='claude-sonnet-4-6', max_tokens=600, system=system, messages=[{'role':'user','content':user}])
        return response.content[0].text.strip()
    except Exception as error: print(f'Error: Anthropic API call failed: {error}', file=sys.stderr); sys.exit(1)

def save_docx(text: str, path: Path, data: CVData) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document(); setup_document_margins(doc)
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; apply_font(p.add_run(data.contact.name), 18, INK, True)
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; apply_font(p.add_run(data.contact.title), 10, MUTED)
        for paragraph in [part.strip() for part in text.split('\n\n') if part.strip()]:
            p=doc.add_paragraph(); p.paragraph_format.space_after=__import__('docx').shared.Pt(10); p.paragraph_format.line_spacing=1.15; apply_font(p.add_run(paragraph), 11, INK)
        footer=doc.sections[0].footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; apply_font(footer.add_run(f'{data.contact.email} | {data.contact.phone} | {data.contact.github}'), 8, MUTED)
        doc.save(path)
    except Exception as error: print(f'Error: Could not write {path}: {error}', file=sys.stderr)

def save_txt(text: str, path: Path, data: CVData) -> None:
    try: path.write_text(f'{data.contact.name}\n{data.contact.title}\n{data.contact.email}\n{data.contact.phone}\n\n{text}\n', encoding='utf-8')
    except OSError as error: print(f'Error: Could not write {path}: {error}', file=sys.stderr)

def generate(args: argparse.Namespace) -> None:
    job_desc=load_job_description(args); letter=call_anthropic(SYSTEM_PROMPT, build_user_prompt(job_desc, CV_DATA)); output=Path('output'); output.mkdir(exist_ok=True); stem=args.out or f'Cover-Letter-{date.today().isoformat()}'
    if args.format in ('docx','both'):
        path=output/f'{stem}.docx'; save_docx(letter,path,CV_DATA)
        if path.exists(): print(f'Saved → {path}')
    if args.format in ('txt','both'):
        path=output/f'{stem}.txt'; save_txt(letter,path,CV_DATA)
        if path.exists(): print(f'Saved → {path}')

def main():
    parser=argparse.ArgumentParser(description='Generate a tailored cover letter.'); group=parser.add_mutually_exclusive_group(required=True); group.add_argument('--job'); group.add_argument('--file'); parser.add_argument('--format', choices=('docx','txt','both'), default='both'); parser.add_argument('--out'); generate(parser.parse_args())
if __name__ == '__main__': main()
