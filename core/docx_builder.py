"""Reusable, restrained DOCX styling helpers."""
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BLUE, INK, MUTED = RGBColor(37, 99, 235), RGBColor(15, 23, 42), RGBColor(71, 85, 105)
def apply_font(run, size=9, color=MUTED, bold=False):
    run.font.name = 'Arial'; run._element.rPr.rFonts.set(qn('w:ascii'), 'Arial'); run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial'); run.font.size = Pt(size); run.font.color.rgb = color; run.bold = bold
def shade_cell(cell, hex_fill):
    props = cell._tc.get_or_add_tcPr(); shade = OxmlElement('w:shd'); shade.set(qn('w:fill'), hex_fill); props.append(shade)
def add_section_heading(doc, title):
    paragraph = doc.add_paragraph(); paragraph.paragraph_format.space_before = Pt(10); paragraph.paragraph_format.space_after = Pt(4); apply_font(paragraph.add_run(title.upper()), 8, BLUE, True)
def add_text_line(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph(); paragraph.paragraph_format.space_after = Pt(4); paragraph.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix): apply_font(paragraph.add_run(bold_prefix), 9, INK, True); apply_font(paragraph.add_run(text[len(bold_prefix):]))
    else: apply_font(paragraph.add_run(text))
def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style='List Bullet'); paragraph.paragraph_format.space_after = Pt(2); apply_font(paragraph.add_run(item), 9)
def setup_document_margins(doc, top=.55, bottom=.55, left=.62, right=.62):
    section = doc.sections[0]; section.top_margin=Inches(top); section.bottom_margin=Inches(bottom); section.left_margin=Inches(left); section.right_margin=Inches(right)
