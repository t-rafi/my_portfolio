from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = 'Towhidul-Islam-Rafi-CV.docx'
BLUE, INK, MUTED = RGBColor(37,99,235), RGBColor(15,23,42), RGBColor(71,85,105)
doc = Document(); s = doc.sections[0]
s.top_margin=s.bottom_margin=Inches(.55); s.left_margin=s.right_margin=Inches(.62)
styles=doc.styles; styles['Normal'].font.name='Arial'; styles['Normal'].font.size=Pt(9); styles['Normal'].font.color.rgb=MUTED

def font(run, size=9, color=MUTED, bold=False):
    run.font.name='Arial'; run._element.rPr.rFonts.set(qn('w:ascii'),'Arial'); run._element.rPr.rFonts.set(qn('w:hAnsi'),'Arial'); run.font.size=Pt(size); run.font.color.rgb=color; run.bold=bold
def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def section(title):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(title.upper()); font(r,8,BLUE,True); r.font.letter_spacing = None
def text(txt, bold_prefix=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.08
    if bold_prefix and txt.startswith(bold_prefix): r=p.add_run(bold_prefix); font(r,9,INK,True); r=p.add_run(txt[len(bold_prefix):]); font(r)
    else: r=p.add_run(txt); font(r)
def bullets(items):
    for item in items:
        p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(1); p.paragraph_format.left_indent=Inches(.18); p.paragraph_format.first_line_indent=Inches(-.14); r=p.add_run(item); font(r,8.5)

# Header
t=doc.add_table(rows=1, cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
t.columns[0].width=Inches(4.7); t.columns[1].width=Inches(2.45)
c=t.cell(0,0); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
p=c.paragraphs[0]; r=p.add_run('Towhidul Islam '); font(r,22,INK,True); r=p.add_run('Rafi'); font(r,22,BLUE,True)
p=c.add_paragraph(); r=p.add_run('Junior Software Engineer | ASP.NET Core & ERP Systems'); font(r,10,INK,True)
p=c.add_paragraph(); r=p.add_run('Open to work | iTech Velocity | Dec 2025-Present'); font(r,8,BLUE,True)
c=t.cell(0,1); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; p.add_run().add_picture('img.jpeg', width=Inches(.8), height=Inches(1.0))
for line in ['tirafi29@gmail.com', '+880 1540 400 287', 'github.com/t-rafi', 'linkedin.com/in/t-rafi/', 'Dhaka, Bangladesh']:
    p=c.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=p.add_run(line); font(r,7.5,BLUE if 'com/' in line else MUTED)

section('Professional Summary')
text('Software Engineer contributing to Clarra, a large-scale enterprise ERP platform at iTech Velocity. Experienced across client requirements, ASP.NET Core and C# development, RDLC report design, IIS deployment, and production support within a 20+ developer team.')
stats=doc.add_table(rows=1, cols=3); stats.autofit=False
for i,(n,l) in enumerate([('30+','RDLC Reports Built'),('20+','ERP Module Domains'),('20+','Developer Team')]):
    cell=stats.cell(0,i); shade(cell,'EFF6FF'); p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(n); font(r,15,BLUE,True); p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(l); font(r,7, MUTED)

section('Professional Experience')
text('Junior Executive, Software Development | iTech Velocity Ltd. | Dec 2025 - Present', 'Junior Executive, Software Development')
text('Contributing to Clarra, a proprietary enterprise ERP platform, from client requirement discussions through deployment and production support.')
bullets(['Developed and maintained 30+ RDLC business reports for enterprise clients.', 'Built application features and REST API endpoints with ASP.NET Core, C#, EF Core, and SQL Server.', 'Handled IIS deployment, production troubleshooting, client communication, and developer coordination.'])
text('Web Development Intern | Pinovation Tech Ltd. | Jun 2025 - Dec 2025', 'Web Development Intern')
bullets(['Developed responsive web interfaces using HTML, CSS, Bootstrap, and JavaScript.', 'Worked in a team development environment using Git and GitHub.'])

section('Projects')
for title, body in [('Clarra ERP Platform','Professional proprietary ERP spanning Finance, HCM, Sales, Supply Chain, Warehouse, Manufacturing, and more. Contributed reports, APIs, deployment, and production support.'),('Registration Form - ASP.NET Core MVC','Data-entry application built with ASP.NET Core MVC, C#, Entity Framework Core, SQL Server, and Razor Views.'),('C# Learning Journey','Tracked repository documenting progression from C# fundamentals to object-oriented design.')]:
    text(title + ' - ' + body, title)

section('Technical Skills')
text('Backend: C#, ASP.NET Core, EF Core | Database: SQL Server, EF Migrations, Redis | Reporting: RDLC Reports, SSRS, Report Design | Frontend: HTML5/CSS3, JavaScript, Razor Views | Tools: IIS, Git, GitHub')
section('Education')
text('B.Sc. in Computer Science & Engineering | Presidency University of Bangladesh | 2026 - Present | CGPA: 3.9 / 4.0 | In Progress', 'B.Sc. in Computer Science & Engineering')
section('Continuous Learning')
bullets(['ASP.NET Core documentation and Microsoft Learn.', 'Clean Architecture and layered application design through production codebase exposure.', 'SQL Server query optimization, IIS configuration, and production troubleshooting.'])

footer=s.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=footer.add_run('Towhidul Islam Rafi | t-rafi.github.io'); font(r,7,MUTED)
doc.save(OUT)
